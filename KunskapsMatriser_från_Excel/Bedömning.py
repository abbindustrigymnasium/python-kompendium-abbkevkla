import openpyxl #library för att använda excel
from docx import Document #python docx som används för att skapa word dokument, a.k.a .docx filer
from docx.shared import Pt #Importar en masssa små libraries från python docx som låter oss göra specifika saker
from docx.enum.style import WD_STYLE_TYPE #Låter oss styla, d.v.s. ändra olika estetiska egenskaper såsåom exempelvis mellanrum
from docx.enum.text import WD_COLOR_INDEX #låter oss highligta text i många olika färger
from docx.enum.text import WD_TAB_ALIGNMENT 
import datetime #Ett bibliotek för att kolla datum

for students in range(2,31): #Vi går igenom alla elever och skapar ett word dokument för varje. Eleverna representeras av vilken siffra de har i excel-bladen
    document = Document() #skapar ett dokument

    wb = openpyxl.load_workbook("Bedömningar.xlsx") #beskriver vilken excelfil vi vill använda
    sheet1 = wb.get_sheet_by_name("Centralt innehåll") #Vi definerar de olika flikarna
    sheet2 = wb.get_sheet_by_name("Kunskapskrav")
    sheet3 = wb.get_sheet_by_name("kraven")
    sheet4 = wb.get_sheet_by_name("Info")

    paragraph_format = document.styles['Normal'].paragraph_format #Vi väljer standard stylen
    paragraph_format.space_after
    paragraph_format.space_after = Pt(10) #ändrar avståndet efter varje paragraf för att få det att se lite bättre ut
    paragraph_format.space_after.pt

    section = document.sections[0]
    header = section.header #vi lägger till en header
    paragraph = header.paragraphs[0] #python-docx skapar automatiskt en paragraf när man skapar en header, så istället för att lägga till en ny
                                     # paragraf och få ett mellanrum upptill så redigerar vi den vi redan har
    tid = datetime.datetime.now() #kollar datumet
    paragraph.text = "Jan Illian\t"+" "+ "\t"+tid.strftime("%x") #lägger till text i header, \t betyder att vi gör ett stort mellanrum så att en text hamnar till vänster och den andra till höger
    paragraph.style = document.styles["Header"]

    document.add_heading("Kunskapsmatris Programmering 1", level=0) #Lägger till en rubrik, level=0 gör att det blir en stor rubrik som inefattar hela dokumentet
    
    document.add_heading("Klass: "+sheet4.cell(row=3, column=2).value + "    Namn: "+sheet1.cell(row=students, column=1).value, 2) #En rubrik med namn och klass
    document.add_heading("Centralt Innehåll") #Rubrik
    
    centralt=[ sheet1.cell(row=1, column=i).value for i in range(2,10)] #hämtar alla centrala innehåll och lägger de i en lista
    kravlvl = [] #tom array som ska inehålla alla kunskapskrav
    for k in range(1, 12): #k är radens nummer, det finns 11 rader i flik 3
        kravlvl += [sheet3.cell(row=k, column=i).value for i in range(1,5)] #i är kolumner, det finns 4
        #för varje rad går vi igenom alla kolumner och lägger till rutans innehåll i kravlvl arrayen
    centralgoals=[sheet1.cell(row=students, column=i).value for i in range(2,10)] #kollar vilka mål som eleven uppfyllt
    KnowledgeReq = [0] #skapar en lista för vilka betyg eleven har uppnått, vi lägger till 0 eftersom att på första raden (E,C,A raden) ska vi inte sätta något betyg
    KnowledgeReq +=[sheet2.cell(row=students, column=i).value for i in range(2,12)] #ser vilka nivåer eleven uppnått

    for i in centralt: #går igenom alla värden i centralt som inehåller själva målet, inte elevens resultat
        cent = document.add_paragraph(style='List Bullet') #skriver in målen i dokumentet i punktform
        if centralgoals[0] == "Genomfört": #kollar om eleven uppnått målet
            run = cent.add_run(i) #vi lägger till målet som en run
            run.bold = True #Fet text
            if len(centralgoals) == 1: #Om det bara finns ett objekt kvar i centralgoals är vi klara, så då bryts loopen och vi går vidare
                break
            del centralgoals[0] #varje gång vi går igenom loopen tar vi bort index 0 i centralgoals. Detta gör att vi kommer gå igenom alla värden eftersom att index 0 byts
        else: #om det inte stämmer gör vi samma sak fast gör inte texten fet
            run = cent.add_run(i)
            run.bold = False
            if len(centralgoals) == 1:
                break
            del centralgoals[0]

    document.add_heading("Kunskapskrav") #Titel
    table = document.add_table(rows=11, cols=4) #Skapar ett rutnät med 11 rader och 4 kolumner
    table.autofit = True 
    counter = 0 #variabler för att hålla koll på vilket värde vi befinner oss vid i loopen
    SetGrade = 0

    for row in table.rows: #vi går igenom alla rader i rutnätet
        counter = 0 #nollställer counter
        if len(KnowledgeReq) == 0: #om kunskapskrav är tom bryter vi loopen, då är vi klara
            break
        elif KnowledgeReq[0] == "A": #Beronde på vilket betyg eleven uppnått sätter vi SetGrade till den index som motsvarar nivån  
            SetGrade = 3
        elif KnowledgeReq[0] == "C":
            SetGrade = 2
        elif KnowledgeReq[0] == "E":
            SetGrade = 1
        else: #Om eleven inte uppnått något sätter vi SetGrade till 5, vilket gör att inget ändras då det inte finns något index 5
            SetGrade = 5
        del KnowledgeReq[0] #Tar bort första objektet i KnowledgeReq för att gå igenom ett nytt kunskapskrav sen
    
        for cell in row.cells: #Går igenom alla celler(rutor) i raden 
            pt = cell.paragraphs[0] #Precis som med headern skapar varje cell en paragraf automatiskt, så för att undvika fula mellanrum i tabellen
            p = pt.add_run(kravlvl[0]) #vi lägger till ett kunskapskrav i cellen
            if counter == SetGrade: #Om vi är på rätt index för betyget eleven fått
                p.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN #highlightar texten grön
                p.bold = True #gör den THICC
            counter = counter + 1 #vi räknar vilken index vi är på, 0, 1, 2 eller 3
            del kravlvl[0] #Tar bort index 0
            
    document.save('Students/'+sheet1.cell(row=students, column=1).value+'.docx') #sparar ner dokumentet i en mapp. Dokumentet döps till elevens namn